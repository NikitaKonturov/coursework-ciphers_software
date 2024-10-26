import os.path

from docx import Document # Импорт класса Document в программу

def check_file_path(filePath : str):
    if os.path.exists(filePath):
        print(f'The file {filePath} exists, continuing work...')
    else:
        raise Exception(f'The file {filePath} does not exist')        
    
def txt_to_docx(txtFile : str, docxFile : str): # Определение будущей функции
    
    check_file_path(filePath)
    
    doc = Document()
    
    with open(txtFile, 'r', encoding='utf-8') as file: # With гарантирует, что файл автоматически закроется
        # 'r' открывает файл в режиме для чтения, encoding указывает, что текст должен читаться именно в указанной кодировке
        for line in file:
            
            everyNewLine = line.strip()
            
            if everyNewLine == '':
                doc.add_paragraph()
                doc.save('D:/Git/coursework-ciphers_software/file_converters/outputu.docx')
            
            else:
                doc.add_paragraph(line.rstrip())
                doc.save('D:/Git/coursework-ciphers_software/file_converters/outputu.docx')
                    

filePath = r'D:/VsCode/PythonProjects/test.txt'

txt_to_docx(filePath, 'D:/Git/coursework-ciphers_software/file_converters/outputu.docx')
