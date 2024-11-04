import os
import re

from docx import Document

def check_file_path(filePath : str):
    if os.path.exists(filePath):
        print(f'The file {filePath} exists, continuing work...')
    else:
        raise Exception(f'The file {filePath} does not exist')
        
def docx_to_txt(language : str, docxFile : str, txtFile):
    
    check_file_path(filePath)
    
    doc = Document(docxFile)
    
    if (language.lower() == 'ru'):
        with open(txtFile, 'w', encoding='utf-8') as txt_file:
            for para in doc.paragraphs:
                checkPart = ''.join(re.findall(r'[A-Za-z]', para.text))
                cleanedText = ''.join(re.findall(r'[А-Яа-я]', para.text)).upper()
            
                if checkPart:
                    raise Exception(f'The file {filePath} contains symbols from the other language')
            
                if cleanedText:
                    txt_file.write(cleanedText)
    
    elif (language.lower() == 'eng'):
        with open(txtFile, 'w', encoding='utf-8') as txt_file:
            for para in doc.paragraphs:
                checkPart = ''.join(re.findall(r'[А-Яа-я]', para.text))
                cleanedText = ''.join(re.findall(r'[A-Za-z]', para.text)).upper()
            
                if checkPart:
                    raise Exception(f'The file {filePath} contains symbols from the other language')
            
                if cleanedText:
                    txt_file.write(cleanedText)
    
    else:
        raise Exception('You have chosen wrong language')

filePath = r'D:/download/testi.docx'

docx_to_txt('eng', filePath, 'output.txt')
