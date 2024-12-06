import os
from typing import BinaryIO
from pathlib import Path
import re

class FileLanguageError(Exception):
    def __init__(self, errorLanguage: str, message: str = "Invalid language in file..."):
        self.errorLanguage = errorLanguage
        self.message = message
        super().__init__(message)

    def __str__(self):
        return f"{self.message} (Language: {self.errorLanguage})"


from docx import Document

def check_file_path(filePath : str):
    if not os.path.exists(filePath):
        raise FileExistsError(f'The file {filePath} does not exist')
        
def save_open_text_docx_as_txt(language : str, openTextFile : BinaryIO, saveOpenTextTxtFile: Path):
    
    doc = Document(openTextFile)
    
    if (language.lower() == 'ru'):
        with open(saveOpenTextTxtFile, 'w', encoding='utf-8') as txtFile:
            for para in doc.paragraphs:
                checkPart = ''.join(re.findall(r'[A-Za-z]', para.text))
                cleanedText = ''.join(re.findall(r'[А-Яа-я]', para.text)).upper()
            
                if checkPart:
                    raise FileLanguageError(f'The file contains symbols from the other language', errorLanguage="en")
            
                if cleanedText:
                    txtFile.write(cleanedText)
    
    elif (language.lower() == 'en'):
        with open(saveOpenTextTxtFile, 'w', encoding='utf-8') as txtFile:
            for para in doc.paragraphs:
                checkPart = ''.join(re.findall(r'[А-Яа-я]', para.text))
                cleanedText = ''.join(re.findall(r'[A-Za-z]', para.text)).upper()
            
                if checkPart:
                    raise FileLanguageError(f'The file contains symbols from the other language', errorLanguage="ru")
            
                if cleanedText:
                    txtFile.write(cleanedText)
    
    else:
        raise FileLanguageError(errorLanguage= "Anny", message='You have chosen wrong language')

def save_docx_as_txt(textFile : BinaryIO, saveTxtFile: Path):
    doc = Document(textFile)
    with open(saveTxtFile, "w", encoding = 'utf-8') as txtFile:
        for paragraph in doc.paragraphs:
             txtFile.write(paragraph.text)
    
    return 
