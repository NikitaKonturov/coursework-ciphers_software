import os.path
from pathlib import Path

from docx import Document # Импорт класса Document в программу

def check_file_path(filePath : str):
    if os.path.exists(filePath):
        print(f'The file {filePath} exists, continuing work...')
    else:
        raise Exception(f'The file {filePath} does not exist')        
    
def save_to_docx(data: dict[str, str], docxFile : Path):
    check_file_path(docxFile.parent)
    doc = Document()    
    for key in data:
        doc.add_paragraph(key)
        doc.add_paragraph(data[key])
        doc.save(docxFile.__str__())
    return 
